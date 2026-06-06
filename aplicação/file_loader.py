"""
Componente 1: FileLoader
Responsabilidade única: ler documentos (.docx, .pdf) a partir dos caminhos
guardados no MySQL e devolver uma lista de LangChain Documents.
"""

import os
import pdfplumber
from pathlib import Path
from docx import Document as DocxDocument
from langchain_core.documents import Document
import mysql.connector
from db_config import DB_CONFIG


BASE_DIR  = Path(os.path.abspath(__file__)).parent   # .../Gest.AI/Aplicação
ROOT_DIR  = BASE_DIR.parent                           # .../Gest.AI


def _resolver_caminho(caminho_db: str) -> str | None:
    """
    Resolve o caminho do ficheiro no disco a partir do caminho relativo
    guardado na BD. Tenta várias bases e, como fallback, faz uma pesquisa
    fuzzy pelo nome do ficheiro dentro de ROOT_DIR.
    """
    # Normaliza separadores (a BD guarda com '/', Windows usa '\')
    caminho_norm = Path(caminho_db)

    # Tentativa 1 e 2: join direto com ROOT_DIR e BASE_DIR
    for base in [ROOT_DIR, BASE_DIR]:
        candidato = base / caminho_norm
        if candidato.exists():
            return str(candidato)

    # Tentativa 3: fallback fuzzy — procura o ficheiro pelo nome dentro de ROOT_DIR
    # Útil quando o caminho na BD está ligeiramente errado mas o ficheiro existe
    nome_ficheiro = caminho_norm.name
    for ficheiro in ROOT_DIR.rglob(nome_ficheiro):
        # Verifica se pelo menos a pasta pai bate certo (evita falsos positivos
        # quando há vários ficheiros com o mesmo nome, ex: "Ficha de serviço atribuido.docx")
        pasta_bd   = caminho_norm.parent.name.lower()
        pasta_real = ficheiro.parent.name.lower()
        if pasta_bd == pasta_real:
            return str(ficheiro)

    return None


def _extrair_nome_pasta(caminho_obj: Path) -> str:
    """Extrai o nome da pasta 'Contrato_*' a partir do caminho do ficheiro."""
    for parte in caminho_obj.parts:
        if parte.startswith("Contrato_"):
            return parte
    return "Geral"


def _ler_docx(caminho: str, metadata: dict) -> Document | None:
    doc = DocxDocument(caminho)
    texto = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if texto.strip():
        return Document(page_content=texto, metadata={**metadata, "tipo": "docx"})
    return None


def _ler_pdf(caminho: str, metadata: dict) -> list[Document]:
    docs = []
    with pdfplumber.open(caminho) as pdf:
        for num_p, page in enumerate(pdf.pages):
            texto = page.extract_text() or ""
            if texto.strip():
                docs.append(Document(
                    page_content=texto,
                    metadata={**metadata, "tipo": "pdf", "page": num_p + 1}
                ))
    return docs


def carregar_documentos() -> list[Document]:
    """
    Liga ao MySQL, obtém os registos da tabela `documentos`
    e devolve uma lista de LangChain Documents prontos a usar.
    """
    documentos = []

    try:
        conn = mysql.connector.connect(**DB_CONFIG)
        cursor = conn.cursor(dictionary=True)
        cursor.execute("SELECT nome, caminho, categoria FROM documentos")
        registos = cursor.fetchall()
        cursor.close()
        conn.close()
    except Exception as e:
        print(f"[FileLoader] Erro MySQL: {e}")
        return []

    print(f"[FileLoader] {len(registos)} documentos encontrados na BD.")

    for reg in registos:
        caminho = _resolver_caminho(reg["caminho"])
        if not caminho:
            print(f"[FileLoader] Aviso: '{reg['nome']}' não encontrado.")
            continue

        caminho_obj = Path(caminho)
        metadata = {
            "source":    reg["nome"],
            "categoria": reg["categoria"],
            "pasta":     _extrair_nome_pasta(caminho_obj),
            "periodo":   caminho_obj.parent.name,
        }

        try:
            ext = caminho.lower()
            if ext.endswith(".docx"):
                doc = _ler_docx(caminho, metadata)
                if doc:
                    documentos.append(doc)
            elif ext.endswith(".pdf"):
                documentos.extend(_ler_pdf(caminho, metadata))
        except Exception as e:
            print(f"[FileLoader] Erro ao processar {reg['nome']}: {e}")

    print(f"[FileLoader] {len(documentos)} documentos carregados.")
    return documentos
