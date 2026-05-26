import docx
import os
import re
import json
import mysql.connector
from db_config import DB_CONFIG

# ---------------------------------------------------------
# NOVIDADE: MAPEAMENTO PARA A CATEGORIA DA BASE DE DADOS
# ---------------------------------------------------------
# Associa o que vem do teu JSON à categoria exata que tens na tua BD
CATEGORIA_POR_TIPO = {
    "integral": "tempo integral anual",
    "parcial": "tempo parcial semestral",
    "parcial-edital": "tempo parcial edital"
}

# Pasta principal onde os documentos preenchidos vão ser guardados
pasta_output = os.path.normpath(os.path.join(os.path.dirname(os.path.abspath(__file__)), "..", "Modelos Contratuais", "Modelos Gerados"))

def criar_nome_pasta_limpo(nome_completo):
    """
    Pega no nome do docente e transforma num nome de pasta válido e limpo.
    """
    nome = nome_completo.replace("Professor ", "").replace("Professora ", "")
    nome = nome.replace("Prof. Dr. ", "").replace("Prof. ", "")
    nome = nome.strip().replace(" ", "_")
    nome_limpo = re.sub(r'[\\/*?:"<>|]', "", nome)
    return f"Contrato_{nome_limpo}"

def processar_renovacao(dados_contrato):
    # 1. Determinar o tipo de contrato e a categoria correspondente na BD
    tipo_contrato = dados_contrato.get("tipo_contrato", "integral").lower()
    categoria_bd = CATEGORIA_POR_TIPO.get(tipo_contrato)
    
    if not categoria_bd:
        print(f"[!] Erro: O tipo '{tipo_contrato}' não tem uma categoria associada no sistema.")
        return

    print(f"[*] A processar contratação do tipo: {tipo_contrato.upper()} (Categoria BD: '{categoria_bd}')")

    # 2. Criar a subpasta do docente
    # [ALTERAÇÃO] Retiradas as chaves da pesquisa pelo nome, visto que o HTML envia agora "nome_docente"
    nome_docente = dados_contrato.get("nome_docente", "Docente_Desconhecido")
    nome_subpasta = criar_nome_pasta_limpo(nome_docente)
    caminho_final = os.path.join(pasta_output, nome_subpasta)
    
    if not os.path.exists(caminho_final):
        os.makedirs(caminho_final)
        print(f"[*] Criada pasta para o processo: {caminho_final}")

    # 3. Ligar à BD e procurar TODOS os templates dessa categoria!
    templates_encontrados = []
    try:
        ligacao = mysql.connector.connect(**DB_CONFIG)
        cursor = ligacao.cursor(dictionary=True)
        
        # MAGIA AQUI: Em vez de procurar por nomes, procuramos por categoria!
        query = "SELECT nome, caminho FROM documentos WHERE categoria = %s"
        cursor.execute(query, (categoria_bd,))
        templates_encontrados = cursor.fetchall()
        cursor.close()
        ligacao.close()
        
    except Exception as e:
        print(f"[!] Erro ao ligar à base de dados MySQL: {e}")
        return

    if not templates_encontrados:
        print(f"[!] Aviso: Nenhum documento encontrado na BD para a categoria '{categoria_bd}'.")
        return

    print(f"[*] Encontrados {len(templates_encontrados)} templates na categoria '{categoria_bd}'. A preencher...")

    # 4. Processar cada template encontrado
    for template in templates_encontrados:
        nome_ficheiro = template["nome"]
        caminho_template = template["caminho"]

        if not os.path.exists(caminho_template):
            print(f"[Aviso] O ficheiro '{nome_ficheiro}' está na BD, mas não foi encontrado no disco: {caminho_template}")
            continue

        try:
            documento = docx.Document(caminho_template)

            # Criamos um dicionário de substituição limpo
            # Isto garante que 'nome_docente' e '{{nome_docente}}' funcionem
            mapa_substituicao = {}
            for k, v in dados_contrato.items():
                chave_limpa = k.replace("{", "").replace("}", "").strip()
                mapa_substituicao[f"{{{chave_limpa}}}"] = str(v)

            # Função auxiliar para substituir texto em parágrafos e tabelas
            def substituir_texto(container):
                for p in container:
                    for tag, valor in mapa_substituicao.items():
                        if tag in p.text:
                            # [CORREÇÃO CRÍTICA] Fazemos a substituição direta
                            p.text = p.text.replace(tag, valor)

            # Aplicar nos parágrafos
            substituir_texto(documento.paragraphs)

            # Aplicar em todas as tabelas
            for tabela in documento.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        substituir_texto(celula.paragraphs)

            # Guardar o ficheiro
            caminho_output_doc = os.path.join(caminho_final, nome_ficheiro)
            documento.save(caminho_output_doc)
            print(f"[*] Sucesso: Ficheiro '{nome_ficheiro}' gerado.")
            
        except Exception as e:
            print(f"[!] Erro ao processar o documento '{nome_ficheiro}': {e}")
            
    print("\n[*] Processo concluído com sucesso!")


# Executar a função com dados de teste 
if __name__ == '__main__':
    caminho_json = os.path.join(os.path.dirname(os.path.abspath(__file__)), "testeDinamicoTemplates.json")
    
    try:
        with open(caminho_json, 'r', encoding='utf-8') as ficheiro:
            dados_externos = json.load(ficheiro)
            
        print("[*] Dados carregados do ficheiro JSON.")
        processar_renovacao(dados_externos)
        
    except FileNotFoundError:
        print(f"[!] Erro: Ficheiro '{caminho_json}' não encontrado.")