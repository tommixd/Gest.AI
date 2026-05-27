import os
import torch
import warnings
import pdfplumber
import mysql.connector
import re
from pathlib import Path
from db_config import DB_CONFIG
from collections import defaultdict
from langchain_core.documents import Document
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from docx import Document as DocxDocument

# NOVO IMPORT PARA O MOTOR GGUF DA NVIDIA/RTX 4060
from llama_cpp import Llama

warnings.filterwarnings("ignore")
dispositivo = "cuda" if torch.cuda.is_available() else "cpu"

# Definir caminhos base
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = os.path.dirname(BASE_DIR)

# --- CONFIGURAÇÃO DA BASE DE DADOS MYSQL ---
DB_CONFIG = {
    'host': 'localhost',
    'user': 'root',            
    'password': '04072002Tomas!', 
    'database': 'BaseDadosGestAI' 
}

# Variável global para mapear pasta -> lista de chunks (usado na resposta)
PASTA_TO_CHUNKS = defaultdict(list)

def ler_documentos_via_mysql():
    """Liga-se ao MySQL e obtém documentos, extraindo o nome da pasta."""
    
    documentos_langchain = []
    
    try:
        ligacao = mysql.connector.connect(**DB_CONFIG)
        cursor = ligacao.cursor(dictionary=True)
        cursor.execute("SELECT nome, caminho, categoria FROM documentos")
        registos = cursor.fetchall()
        cursor.close()
        ligacao.close()
    except Exception as e:
        print(f"[!] Erro MySQL: {e}")
        return []

    print(f"[*] Encontrados {len(registos)} documentos na BD.")

    for registo in registos:
        caminho_db = registo["caminho"]
        nome_doc = registo["nome"]
        categoria = registo["categoria"]
        
        caminho_ficheiro = os.path.join(ROOT_DIR, caminho_db).replace('\\', '/')
        if not os.path.exists(caminho_ficheiro):
            caminho_ficheiro = os.path.join(BASE_DIR, caminho_db).replace('\\', '/')
            if not os.path.exists(caminho_ficheiro):
                print(f"[Aviso] '{nome_doc}' não encontrado em: {caminho_ficheiro}")
                continue

        caminho_obj = Path(caminho_ficheiro)
        # Estrutura: Contrato_Name/2025-2026/ficheiro.docx
        # parent = 2025-2026, parent.parent = Contrato_Name
        periodo = caminho_obj.parent.name
        nome_pasta = "Geral"
        #Procura o nome da pasta que começa com "Contrato_" e tem o nome do docente
        for parte in caminho_obj.parts:
            if parte.startswith("Contrato_"):
                nome_pasta = parte
                break

            
        try:
            if caminho_ficheiro.lower().endswith(".docx"):
                doc_word = DocxDocument(caminho_ficheiro)
                texto = "\n".join([p.text for p in doc_word.paragraphs if p.text.strip()])
                if texto.strip():
                    documentos_langchain.append(Document(
                        page_content=texto, 
                        metadata={
                            "source": nome_doc,
                            "categoria": categoria,
                            "tipo": "docx",
                            "pasta": nome_pasta,
                            "periodo": periodo
                        }
                    ))

            elif caminho_ficheiro.lower().endswith(".pdf"):
                with pdfplumber.open(caminho_ficheiro) as pdf:
                    for num_p, page in enumerate(pdf.pages):
                        texto = page.extract_text() or ""
                        if texto.strip():
                            documentos_langchain.append(Document(
                                page_content=texto, 
                                metadata={
                                    "source": nome_doc,
                                    "categoria": categoria,
                                    "page": num_p + 1,
                                    "tipo": "pdf",
                                    "pasta": nome_pasta,
                                    "periodo": periodo
                                }
                            ))
        except Exception as e:
            print(f"[!] Erro ao processar {nome_doc}: {e}")

    print(f"[*] Documentos carregados: {len(documentos_langchain)}")
    return documentos_langchain


def obter_tabelas_com_coluna_nome_docente():
    """Retorna as tabelas que contêm a coluna nome_docente na BD atual."""
    try:
        db = mysql.connector.connect(**DB_CONFIG)
        cursor = db.cursor()
        cursor.execute(
            """
            SELECT table_name
            FROM information_schema.columns
            WHERE table_schema = %s
              AND column_name = 'nome_docente'
            """,
            (DB_CONFIG['database'],)
        )
        tabelas = [row[0] for row in cursor.fetchall()]
        cursor.close()
        db.close()
        return tabelas
    except Exception as e:
        print(f"[!] Erro ao obter tabelas com nome_docente: {e}")
        return []


def buscar_registos_por_nome(nome_pessoa):
    """Procura nas tabelas 'docentes' e 'rascunhos' pelos registos do docente."""
    resultados = []
    
    try:
        db = mysql.connector.connect(**DB_CONFIG)
        cursor = db.cursor(dictionary=True)
        
        # 1. Procurar na tabela DOCENTES (onde a coluna se chama 'nome')
        try:
            cursor.execute(
                "SELECT * FROM docentes WHERE nome LIKE %s LIMIT 5",
                (f"%{nome_pessoa}%",)
            )
            for registo in cursor.fetchall():
                resultados.append({
                    'tabela': 'docentes',
                    'registo': registo
                })
        except Exception as e:
            print(f"[!] Erro ao consultar docentes: {e}")

        # 2. Procurar na tabela RASCUNHOS (onde a coluna se chama 'nome_docente')
        try:
            cursor.execute(
                "SELECT * FROM rascunhos WHERE nome_docente LIKE %s LIMIT 5",
                (f"%{nome_pessoa}%",)
            )
            for registo in cursor.fetchall():
                resultados.append({
                    'tabela': 'rascunhos',
                    'registo': registo
                })
        except Exception as e:
            print(f"[!] Erro ao consultar rascunhos: {e}")
            
        cursor.close()
        db.close()
    except Exception as e:
        print(f"[!] Erro ao conectar à BD para buscar registos: {e}")

    return resultados


def formatar_registo_bd(tabela, registo):
    """Formata um registo para texto compacto no prompt."""
    campos = []
    for chave, valor in registo.items():
        if valor is None:
            continue
        campos.append(f"{chave}={valor}")
        if len(campos) >= 6:
            break
    return f"[Registo do Sistema Central - Tabela: {tabela}] " + ", ".join(campos)


def inicializar_rag():
    """Inicializa o sistema RAG e cria o mapeamento pasta->chunks."""
    global PASTA_TO_CHUNKS
    print(f"[*] Motor: {dispositivo.upper()}")
    
    docs = ler_documentos_via_mysql()
    if not docs:
        print("[!] Nenhum documento carregado.")
        return None, None

    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=300,
        separators=["\n\n", "\n", ".", " ", ""]
    )
    chunks = splitter.split_documents(docs)

    PASTA_TO_CHUNKS.clear()
    for chunk in chunks:
        pasta = chunk.metadata.get('pasta', 'SEM_PASTA')
        PASTA_TO_CHUNKS[pasta].append(chunk)

    print("[*] A carregar embeddings no CPU para poupar VRAM...")
    embeddings = HuggingFaceEmbeddings(
        model_name="sentence-transformers/paraphrase-multilingual-MiniLM-L12-v2",
        model_kwargs={'device': 'cpu'}
    )
    
    db = FAISS.from_documents(chunks, embeddings)
    
    # --- NOVO BLOCO GGUF PARA A RTX 4060 ---
    print("[*] A carregar Qwen2.5-Coder-7B GGUF...")
    model_path = os.path.join(BASE_DIR, "Qwen2.5.1-Coder-7B-Instruct-Q4_K_M.gguf")
    llm = Llama(
        model_path=model_path, # Caminho absoluto baseado no diretório do script
        n_gpu_layers=-1, 
        n_ctx=8192,      
        verbose=True    
    )
    
    # Aumentado o número de documentos recuperados (k) para dar mais contexto
    retriever = db.as_retriever(search_kwargs={"k": 5, "fetch_k": 20})
    return retriever, llm


def responder_pergunta(pergunta, retriever, llm):
    global PASTA_TO_CHUNKS
    
    nome_pessoa = None
    # Melhorada a regex para extração de nomes (mais flexível)
    match = re.search(r'(?:da|de|do|d[ae])?\s*([A-ZÀ-Ÿ][a-zà-ÿ]+(?:\s+[A-ZÀ-Ÿ][a-zà-ÿ]+)+)', pergunta)
    if match:
        nome_pessoa = match.group(1).strip()

    # Procura também por anos na pergunta (ex: 2025, 2026/2027)
    anos_na_pergunta = re.findall(r'\b(20\d{2}(?:/\d{2,4})?)\b', pergunta)

    # se escrever com minusculas, o sistema procura o nome na lista de pastas que já conhece
    if not nome_pessoa:
        pergunta_lower = pergunta.lower()
        for pasta in PASTA_TO_CHUNKS.keys():
            if pasta.startswith("Contrato_"):
                nome_docente_pasta = pasta.replace("Contrato_", "").replace("_", " ").lower()
                if nome_docente_pasta and nome_docente_pasta in pergunta_lower:
                    nome_pessoa = pasta.replace("Contrato_", "").replace("_", " ")
                    break

    print(f"[DEBUG] Nome extraído: {nome_pessoa}")
    print(f"[DEBUG] Anos extraídos: {anos_na_pergunta}")
    
    info_bd = ""
    if nome_pessoa:
        registos_bd = buscar_registos_por_nome(nome_pessoa)
        if registos_bd:
            info_partes = []
            for item in registos_bd:
                texto_registo = formatar_registo_bd(item['tabela'], item['registo'])
                info_partes.append(texto_registo)
                if item['tabela'] == 'rascunhos':
                    rascunho = item['registo']
                    info_partes.append(
                        f"O processo de {rascunho.get('nome_docente')} encontra-se guardado como um RASCUNHO (Incompleto) do tipo '{rascunho.get('tipo_contrato')}'."
                    )
                    print(f"[DEBUG] Rascunho encontrado para: {nome_pessoa}")

            info_bd = "\n".join(info_partes)
        else:
            print(f"[DEBUG] Nenhum registo encontrado em tabelas com nome_docente para: {nome_pessoa}")

    docs_rel = retriever.invoke(pergunta)
    
    pasta_forcada = None
    if nome_pessoa:
        # Normalização do nome para procura de pasta
        nome_normalizado = nome_pessoa.replace(' ', '_').lower()
        for pasta in PASTA_TO_CHUNKS.keys():
            if nome_normalizado in pasta.lower():
                pasta_forcada = pasta
                break
        
        if pasta_forcada:
            chunks_adicionais = PASTA_TO_CHUNKS[pasta_forcada]
            existing = {(d.page_content, d.metadata.get('source')) for d in docs_rel}
            for chunk in chunks_adicionais:
                key = (chunk.page_content, chunk.metadata.get('source'))
                if key not in existing:
                    docs_rel.append(chunk)
                    existing.add(key)
    
    if not docs_rel and not info_bd:
        return "Não encontrei documentos ou registos na base de dados sobre isso.", ""
    
    # Ordenação inteligente: Prioriza pasta forçada E depois ordena por período (mais recente primeiro ou por relevância)
    if pasta_forcada:
        # Coloca documentos da pasta forçada no topo, mas mantém a diversidade de períodos
        docs_rel = sorted(docs_rel, key=lambda d: (
            0 if d.metadata.get('pasta') == pasta_forcada else 1,
            -int(re.search(r'\d+', d.metadata.get('periodo', '0')).group()) if re.search(r'\d+', d.metadata.get('periodo', '')) else 0
        ))
    
    # Aumentado o limite de documentos para o contexto da LLM para capturar múltiplos períodos
    docs_rel = docs_rel[:10]
    
    contexto_parts = []
    for i, d in enumerate(docs_rel, 1):
        fonte = d.metadata.get('source', '?')
        pasta = d.metadata.get('pasta', 'Geral')
        periodo = d.metadata.get('periodo', 'Desconhecido')
        conteudo = d.page_content
        contexto_parts.append(f"[Ficheiro: {fonte} | Pasta: {pasta} | Período: {periodo}]\n{conteudo}\n")
    contexto = "\n".join(contexto_parts)
    
    conhecimento_sistema = ""
    if info_bd:
        conhecimento_sistema += f"Informação atual da base de dados:\n{info_bd}\n\n"
    if contexto:
        conhecimento_sistema += f"Informação extraída dos Documentos do processo:\n{contexto}\n\n"
    if not conhecimento_sistema:
        conhecimento_sistema = "Não possuis qualquer informação no momento sobre este assunto."

    instrucao_adicional = ""
    if pasta_forcada:
        instrucao_adicional = f"Foca-te prioritariamente na informação da pasta '{pasta_forcada}' para responder a esta pergunta."
    
    if anos_na_pergunta:
        instrucao_adicional += f" Presta especial atenção aos dados referentes ao(s) ano(s)/período(s): {', '.join(anos_na_pergunta)}."

    prompt = f"""<|im_start|>system
És o Gest.AI, um assistente inteligente de recursos humanos e gestão académica.
A tua missão é responder à pergunta do utilizador de forma natural, direta e prestável.

Abaixo está a tua base de conhecimento atual sobre este caso. Usa APENAS esta informação para responder. 

REGRA DE OURO: Age como se este conhecimento fosse teu. Não uses frases robóticas como "De acordo com os documentos" ou "A base de dados diz", EXCETO se o utilizador te perguntar especificamente "De onde tiraste essa informação?", "Qual é o ficheiro?" ou "Qual a fonte?". Nesses casos, podes referir os nomes dos ficheiros e pastas indicados no conhecimento.

REGRA 1: Se a resposta não estiver no contexto, responde "Não encontrei informação sobre isso nos documentos." NUNCA inventes dados.
REGRA 2: Sempre que justificares uma resposta com base nos documentos, DEVES citar o nome do ficheiro que usaste (ex: "Segundo o ficheiro X.docx...").
REGRA 3: Sê direto e profissional.
REGRA 4: Se a informação vier de um "[Registo do Sistema Central...]", deves dizer que verificaste na "base de dados do sistema" ou na "ficha de docente" e NUNCA num ficheiro.
REGRA 5: Se a informação vier de um "[Ficheiro: ...]", podes e deves referir o nome desse documento/ficheiro.

REGRA DE HISTÓRICO: O utilizador valoriza muito a precisão cronológica. Se houver dados diferentes para anos letivos diferentes (ex: 2025/2026 vs 2026/2027), DEVES distinguir claramente as situações na tua resposta. Se o utilizador não especificar o ano, apresenta a informação de forma estruturada por período.

[INÍCIO DO CONHECIMENTO]
{conhecimento_sistema}
[FIM DO CONHECIMENTO]

{instrucao_adicional}
<|im_end|>
<|im_start|>user
{pergunta}
<|im_end|>
<|im_start|>assistant
"""
    
    try:
        resposta_raw = llm(
            prompt,
            max_tokens=768,       # Aumentado para permitir respostas mais detalhadas sobre múltiplos períodos
            temperature=0.1,
            stop=["<|im_end|>"]
        )
        resposta = resposta_raw["choices"][0]["text"].strip()
    except Exception as e:
        return f"Erro ao gerar resposta da IA: {e}", ""
    
    # Adição de fontes usadas de forma mais abrangente
    if pasta_forcada:
        docs_da_pasta = [d for d in docs_rel if d.metadata.get('pasta') == pasta_forcada]
        if docs_da_pasta:
            fontes_usadas = list(set([d.metadata.get('source', '?') for d in docs_da_pasta]))
            fontes_str = ", ".join(fontes_usadas)
            if not f"*(Fonte: {fontes_str})*" in resposta:
                resposta += f"\n\n*(Fonte: {fontes_str})*"
            
    return resposta, contexto