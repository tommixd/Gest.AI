import os
import sys
import threading
import re
import docx
import json
import shutil
import mysql.connector
from datetime import datetime, date
from pathlib import Path
#vários imports ainda não utilizados mas que serão uteis para futuras funcionalidades (ex: gestão de utilizadores, autenticação, etc.)
from flask import Flask, render_template, g, request, send_file, jsonify, abort
from flask_sqlalchemy import SQLAlchemy
from werkzeug.utils import secure_filename 
from functools import wraps
from db_config import DB_CONFIG

# --- GARANTIR IMPORTAÇÃO CORRETA DO RAG_LLM ---
BASE_DIR = os.path.abspath(os.path.dirname(__file__))
if BASE_DIR not in sys.path:
    sys.path.insert(0, BASE_DIR)

try:
    import RAG_LLM
    import importlib
    importlib.reload(RAG_LLM)
    from RAG_LLM import responder_pergunta, inicializar_rag
    print("[OK] Funções do RAG_LLM carregadas com sucesso.")
except Exception as e:
    print(f"[!] Erro crítico ao carregar RAG_LLM: {e}")
    def responder_pergunta(*args): return "IA em manutenção.", ""
    def inicializar_rag(): return None, None, None

try:
    from templates import processar_renovacao
    print("[OK] Motor de templates importado com sucesso.")
except ImportError:
    print("[!] Aviso: Não foi possível importar templates.py. A usar lógica interna.")
    def processar_renovacao(dados): return None

# Inicializa a aplicação
app = Flask(__name__)
app.config["SITE_NAME"] = "Gest.AI"  


LABELS_HISTORICO = {
    #id de contrato
    "tipo" : "Tipo de Processo",
    "tipo_contrato":         "Tipo de Contrato",
    "nome_docente":          "Nome do Docente",
    "ano_letivo":            "Ano Letivo",
    "ano_anterior":          "Ano Anterior",
    "data_inicio_contrato":  "Data de Início",
    "data_fim_contrato":     "Data de Fim",
    #dados académicos
    "area_contratacao":      "Área de Contratação",
    "areas_curriculares":    "Áreas Curriculares",
    "departamento":          "Departamento",
    "tipo_docente":          "Tipo de Docente",
    "funcoes_externas":      "Funções Externas",
    # Carga horária
    "total_horas_contacto":  "Total de Horas de Contacto",
    "horario_semanal":       "Horário Semanal",
    # Júri
    "profAAA":               "Membro do Júri A",
    "profBBB":               "Membro do Júri B",
    "profCCC":               "Membro do Júri C",
    "profDDD":               "Membro do Júri D",
    "profXXX":               "Presidente do Júri / Diretor",
    # Observações
    "relatorio_juri":        "Relatório / Observações",
}

#chaves internars que não devem aparecer na ui
CHAVES_OCULTAS = {"tipo", "tipo_contrato", "rascunho_id"}

@app.template_filter('historico_legivel')
def historico_legisvel(valor):
    """Converte o json do historico num bloco html legivel para o user"""

    import re
    if isinstance(valor, str):
        try:
            valor = json.loads(valor)
        except Exception:
            return valor # se falhar devolve a string como estava
    if not isinstance(valor, dict):
        return str(valor) # se não for um dict, devolve como estava
    
    linhas = []
    for chave, val in valor.items():
        if not val or chave in CHAVES_OCULTAS:
            continue
        label = LABELS_HISTORICO.get(chave)
        if not label:
            label = re.sub(r'_', ' ', chave).title()
        linhas.append(f'<div class="hist-linha"><span class="hist-label">{label}</span><span class="hist-valor">{val}</span></div>')

    return '\n'.join(linhas)

# Tipos de contratos mostrados na página de documentos
TIPOS_DOCUMENTOS_PERMITIDOS = {
    'tempo integral anual',
    'tempo parcial semestral',
    'tempo parcial edital'
}

# --- CONFIGURAÇÃO DA BASE DE DADOS MYSQL ---
PASTA_UPLOADS = 'PastaUploadsSiteTest'

# --- CONFIGURAÇÃO DA IA EM BACKGROUND ---
global_retriever = None
global_llm = None
global_sql_agent = None
ia_pronta = False

def migrar_contratos_existentes_para_bd():
    """Varre Modelos Gerados e registra contratos antigos que não estão na BD."""
    ROOT_DIR = os.path.dirname(BASE_DIR)
    pasta_modelos = os.path.join(ROOT_DIR, 'Modelos Contratuais', 'Modelos Gerados')
    
    if not os.path.exists(pasta_modelos):
        print("[!] Pasta Modelos Gerados não encontrada.")
        return
    
    print("[*] A varrer contratos existentes no disco...")
    ficheiros_registados = 0
    
    try:
        ligacao = mysql.connector.connect(**DB_CONFIG)
        cursor = ligacao.cursor(dictionary=True)
        
        # Obter lista de documentos já registados
        cursor.execute("SELECT caminho FROM documentos WHERE categoria='gerado'")
        caminhos_existentes = {row["caminho"] for row in cursor.fetchall()}
        
        # Varrer estrutura: Contrato_Name/2025-2026/ficheiro.docx
        for nome_contrato in os.listdir(pasta_modelos):
            pasta_contrato = os.path.join(pasta_modelos, nome_contrato)
            if not os.path.isdir(pasta_contrato):
                continue
            
            for periodo in os.listdir(pasta_contrato):
                pasta_periodo = os.path.join(pasta_contrato, periodo)
                if not os.path.isdir(pasta_periodo):
                    continue
                
                for ficheiro in os.listdir(pasta_periodo):
                    if not ficheiro.lower().endswith(('.docx', '.pdf')):
                        continue
                    
                    # Caminho relativo para a BD
                    caminho_relativo = os.path.join('Modelos Contratuais', 'Modelos Gerados', 
                                                    nome_contrato, periodo, ficheiro).replace('\\', '/')
                    
                    if caminho_relativo in caminhos_existentes:
                        continue
                    
                    try:
                        data_upload = datetime.now().strftime('%Y-%m-%d')
                        cursor.execute(
                            """
                            INSERT INTO documentos (nome, caminho, categoria, data_upload, versao_contrato)
                            VALUES (%s, %s, %s, %s, %s)
                            """,
                            (ficheiro, caminho_relativo, 'gerado', data_upload, 1)
                        )
                        ficheiros_registados += 1
                        print(f"[+] Registado: {caminho_relativo}")
                    except Exception as e:
                        print(f"[!] Erro ao registar {ficheiro}: {e}")
        
        ligacao.commit()
        cursor.close()
        ligacao.close()
        print(f"[OK] Migração concluída: {ficheiros_registados} documentos registados.")
    except Exception as e:
        print(f"[!] Erro na migração: {e}")

def iniciar_ia_background():
    global global_retriever, global_llm, global_sql_agent, ia_pronta
    try:
        print("\n[A inicializar o cérebro Gest.AI em 2º plano...]")
        global_retriever,  global_llm, global_sql_agent = inicializar_rag()
        if global_retriever and global_llm and global_sql_agent:
            ia_pronta = True
            print("\n[=========================================]")
            print("[ IA PRONTA! O Chatbox já pode responder. ]")
            print("[=========================================]\n")
    except Exception as e:
        print(f"\n[!] Erro crítico ao carregar a IA: {e}")

thread_ia = threading.Thread(target=iniciar_ia_background)
thread_ia.daemon = True

# Executar migração de contratos antigos antes de iniciar a IA
print("\n[Verificando contratos existentes...]")
migrar_contratos_existentes_para_bd()

# Depois iniciar a IA
thread_ia.start()

# --- MAPEAMENTO E LÓGICA DE TEMPLATES ---
CATEGORIA_POR_TIPO = {
    "renovacao-integral": "tempo integral anual",
    "renovacao-parcial": "tempo parcial semestral",
    "primeira-vez": "tempo parcial edital"
}

def criar_nome_pasta_limpo(nome_completo):
    nome = nome_completo.replace("Professor ", "").replace("Professora ", "")
    nome = nome.replace("Prof. Dr. ", "").replace("Prof. ", "").replace("Prof. Doutor ", "")
    nome = nome.strip().replace(" ", "_")
    nome_limpo = re.sub(r'[\\/*?:"<>|]', "", nome)
    return f"Contrato_{nome_limpo}"

def obter_ano_letivo_destino(dados_formulario):
    """Extrai o ano letivo do formulário ou gera baseado na data atual."""
    ano_letivo = dados_formulario.get("ano_letivo", "").strip()
    if ano_letivo:
        return ano_letivo
    # Se não estiver no formulário, gera baseado na data (ex: 2025-2026)
    ano_atual = datetime.now().year
    if datetime.now().month >= 9:
        return f"{ano_atual}-{ano_atual + 1}"
    else:
        return f"{ano_atual - 1}-{ano_atual}"


def processar_templates(dados_formulario, tipo_contratacao):
    categoria_bd = CATEGORIA_POR_TIPO.get(tipo_contratacao)
    if not categoria_bd:
        print(f"[!] Erro: Tipo de contratação '{tipo_contratacao}' não mapeado.")
        return None, f"Tipo de contratação '{tipo_contratacao}' não reconhecido."

    nome_docente = dados_formulario.get("nome_docente", "Docente_Desconhecido")
    nome_subpasta_docente = criar_nome_pasta_limpo(nome_docente)
    ano_letivo = obter_ano_letivo_destino(dados_formulario)
    
    ROOT_DIR = os.path.dirname(BASE_DIR)
    # Estrutura: Contrato_Nome/2025-2026/
    caminho_final = os.path.join(ROOT_DIR, 'Modelos Contratuais', 'Modelos Gerados', nome_subpasta_docente, ano_letivo)
    
    print(f"[*] A iniciar processo para: {nome_docente}")
    print(f"[*] Ano letivo: {ano_letivo}")
    print(f"[*] Categoria alvo na BD: {categoria_bd}")
    print(f"[*] Pasta de destino: {caminho_final}")

    if not os.path.exists(caminho_final):
        os.makedirs(caminho_final)
        print(f"[*] Criada pasta de processo: {nome_subpasta_docente}/{ano_letivo}")

    templates_encontrados = []
    try:
        # Ligar ao MySQL
        ligacao = mysql.connector.connect(**DB_CONFIG)
        cursor = ligacao.cursor(dictionary=True)
        
        cursor.execute("SELECT nome, caminho FROM documentos WHERE categoria LIKE %s", (f"%{categoria_bd}%",))
        templates_encontrados = cursor.fetchall()
        
        cursor.close()
        ligacao.close()
        
        print(f"[*] Templates encontrados na BD: {len(templates_encontrados)}")
    except Exception as e:
        print(f"[!] Erro de SQL: {e}")
        return None, f"Erro ao aceder à base de dados: {e}"

    if not templates_encontrados:
        print(f"[!] Aviso: A categoria '{categoria_bd}' não tem documentos registados.")
        return None, f"Nenhum template encontrado para a categoria '{categoria_bd}'."

    ficheiros_gerados = []
    for template in templates_encontrados:
        nome_ficheiro = template["nome"]
        caminho_abs_template = os.path.join(ROOT_DIR, template["caminho"])
        
        if not os.path.exists(caminho_abs_template):
            print(f"[!] Ficheiro não encontrado no disco: {caminho_abs_template}")
            continue

        try:
            print(f"[*] A preencher: {nome_ficheiro}...")
            documento = docx.Document(caminho_abs_template)

            for paragrafo in documento.paragraphs:
                for chave, valor in dados_formulario.items():
                    if chave in paragrafo.text and valor:
                        paragrafo.text = paragrafo.text.replace(chave, str(valor))

            for tabela in documento.tables:
                for linha in tabela.rows:
                    for celula in linha.cells:
                        for paragrafo in celula.paragraphs:
                            for chave, valor in dados_formulario.items():
                                if chave in paragrafo.text and valor:
                                    paragrafo.text = paragrafo.text.replace(chave, str(valor))

            caminho_output = os.path.join(caminho_final, nome_ficheiro)
            documento.save(caminho_output)
            ficheiros_gerados.append(nome_ficheiro)
            print(f"[OK] Gerado: {nome_ficheiro}")

        except Exception as e:
            print(f"[!] Erro ao processar {nome_ficheiro}: {e}")

    if not ficheiros_gerados:
        return None, "Os templates foram encontrados mas não foi possível gerar nenhum ficheiro."


    # Retorna o caminho completo incluindo o período
    return caminho_final, ficheiros_gerados

# --- ROTAS FLASK E BASE DE DADOS ---
def get_db():
    db = getattr(g, '_database', None)
    if db is None:
        db = g._database = mysql.connector.connect(**DB_CONFIG)
    return db

@app.teardown_appcontext
def close_connection(exception):
    db = getattr(g, '_database', None)
    if db is not None:
        db.close()


@app.route('/')  
def index():
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM documentos ORDER BY categoria, caminho")
    lista_documentos = cursor.fetchall()
    cursor.close()
    
    # Filtra apenas os contratos que a instituição realiza
    lista_documentos = [
        doc for doc in lista_documentos
        if doc['categoria'] in TIPOS_DOCUMENTOS_PERMITIDOS
        and 'Modelos Gerados' not in doc['caminho']
    ]

    # Estrutura: documentos_agrupados[categoria][nome_contrato][subpasta_ou_files] = [lista_de_docs]
    documentos_agrupados = {}
    
    for doc in lista_documentos:
        cat = doc['categoria']
        if cat not in documentos_agrupados: documentos_agrupados[cat] = {}
        
        # Lógica para caminhos em 'Modelos Gerados'
        if 'Modelos Gerados' in doc['caminho']:
            partes = doc['caminho'].replace('\\', '/').split('/')
            
            # Encontrar onde começa o "Contrato_"
            idx_contrato = -1
            for i, p in enumerate(partes):
                if p.startswith('Contrato_'):
                    idx_contrato = i
                    break
            
            if idx_contrato != -1:
                nome_contrato = partes[idx_contrato].replace('_', ' ')
                # Subpastas são tudo o que está entre o contrato e o ficheiro
                subpastas = "/".join(partes[idx_contrato+1:-1])
                pasta_display = subpastas if subpastas else "Ficheiros"
                
                if nome_contrato not in documentos_agrupados[cat]:
                    documentos_agrupados[cat][nome_contrato] = {}
                if pasta_display not in documentos_agrupados[cat][nome_contrato]:
                    documentos_agrupados[cat][nome_contrato][pasta_display] = []
                
                documentos_agrupados[cat][nome_contrato][pasta_display].append(doc)
            else:
                # Caso de segurança se não encontrar "Contrato_"
                if 'Geral' not in documentos_agrupados[cat]: documentos_agrupados[cat]['Geral'] = {'Ficheiros': []}
                documentos_agrupados[cat]['Geral']['Ficheiros'].append(doc)
        else:
            # Caso Geral
            if 'Geral' not in documentos_agrupados[cat]: documentos_agrupados[cat]['Geral'] = {'Ficheiros': []}
            documentos_agrupados[cat]['Geral']['Ficheiros'].append(doc)
            
    return render_template('index.html', documentos_agrupados=documentos_agrupados)


@app.route('/contratacao/<tipo>')
def iniciar_contratacao(tipo):
    titulos = {
        'renovacao-integral': 'Renovação de Contrato (Integral)',
        'renovacao-parcial': 'Renovação de Contrato (Parcial)',
        'primeira-vez': 'Primeira Contratação (Edital)'
    }
    
    regimes = {
        'renovacao-integral': 'Tempo Integral',
        'renovacao-parcial': 'Tempo Parcial',
        'primeira-vez': 'Tempo Parcial Edital'
    }

    titulo = titulos.get(tipo, 'Novo Processo')
    regime_nome = regimes.get(tipo, 'Padrão')
    
    # -----------------------------------------------------------------
    # LÓGICA DO RASCUNHO: Verifica se fomos chamados por um rascunho_id
    # -----------------------------------------------------------------
    rascunho_id = request.args.get('rascunho_id')
    dados_preenchidos = {}

    lista_professores = []
    lista_areas = []
    lista_ucs = []
    lista_cursos = []

    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # Carregar Rascunho
        if rascunho_id:
            cursor.execute("SELECT dados_formulario FROM rascunhos WHERE id = %s", (rascunho_id,))
            resultado = cursor.fetchone()
            if resultado:
                dados_preenchidos = json.loads(resultado['dados_formulario'])
        
        # --- NOVA PARTE: Buscar dados dinâmicos ao MySQL ---
        # Usamos try/except individuais para não quebrar a página se uma tabela falhar
        try:
            cursor.execute("SELECT nome FROM docentes ORDER BY nome")
            lista_professores = [row['nome'] for row in cursor.fetchall()]
        except Exception: pass

        try:
            cursor.execute("SELECT nome FROM areas_estudo ORDER BY nome")
            lista_areas = [row['nome'] for row in cursor.fetchall()]
        except Exception: pass

        try:
            cursor.execute("SELECT nome FROM ucs ORDER BY nome")
            lista_ucs = [row['nome'] for row in cursor.fetchall()]
        except Exception: pass

        try:
            cursor.execute("SELECT nome FROM cursos ORDER BY nome")
            lista_cursos = [row['nome'] for row in cursor.fetchall()]
        except Exception: pass

        lista_cargas = []
        try:
            cursor.execute(
                "SELECT id_carga, tempo_contratual, tempo_aulas, tempo_apoio, tempo_preparacao, percentagem FROM carga_horaria ORDER BY id_carga"
            )
            lista_cargas = cursor.fetchall()
        except Exception: pass
        
        cursor.close()
    except Exception as e:
        print(f"[!] Erro ao carregar dados para o formulário: {e}")

    # Passamos as listas para o HTML usar!
    return render_template('contratacao.html', 
                           titulo=titulo, 
                           regime_nome=regime_nome, 
                           tipo=tipo,
                           tipo_contrato=tipo,
                           form_data=dados_preenchidos,
                           professores=lista_professores,
                           areas=lista_areas,
                           ucs=lista_ucs,
                           cursos=lista_cursos,
                           cargas=lista_cargas)

@app.route('/api/juris-disponiveis')
def juris_disponiveis():
    """Devolve docentes com isJuri = 'sim', excluindo o proprio docente"""
    nome_docente = request.args.get('excluir', '').strip()

    try:
        ligacao = mysql.connector.connect(**DB_CONFIG)
        cursor = ligacao.cursor(dictionary=True)

        query = """
          SELECT nome 
          FROM docentes
          WHERE isJuri = 'sim' AND nome != %s
          ORDER BY nome ASC
        """

        cursor.execute(query, (nome_docente,))
        juris = [row['nome'] for row in cursor.fetchall()]
        cursor.close()
        ligacao.close()
        return jsonify({"juris": juris})
    
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/docente-detalhes')
def docente_detalhes():
    nome = request.args.get('nome', '').strip()
    if not nome:
        return jsonify({"erro": "nome ausente"}), 400

    detalhes = {
        "nome_docente": nome,
        "tipo_docente": None,
        "departamento": None,
        "total_horas_contacto": None,
        "horario_semanal": None
    }

    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # 1. Buscar info base do docente e o Histórico Permanente!
        cursor.execute("SELECT id_docente, tipo_docente, departamento, dados_historico FROM docentes WHERE nome = %s", (nome,))
        docente = cursor.fetchone()
        
        if docente:
            detalhes["tipo_docente"] = docente["tipo_docente"]
            detalhes["departamento"] = docente["departamento"]
            
            # Se tiver histórico guardado, carrega TUDO (UCs, Juris, Áreas, etc.)
            if docente.get("dados_historico"):
                historico = json.loads(docente["dados_historico"])
                for chave, valor in historico.items():
                    if valor:
                        detalhes[chave] = valor
            
            # Buscar dados da Carga horária - obter contrato(s) associado(s) ao docente
            cursor.execute(
                "SELECT c.id_contrato, ch.tempo_contratual, ch.tempo_aulas, ch.tempo_apoio, ch.tempo_preparacao, ch.percentagem "
                "FROM contratos c JOIN carga_horaria ch ON c.carga_horaria_id_carga = ch.id_carga "
                "WHERE c.docentes_id_docente = %s ORDER BY c.id_contrato DESC LIMIT 1",
                (docente["id_docente"],)
            )
            contrato = cursor.fetchone()
            if contrato:
                detalhes["total_horas_contacto"] = contrato["tempo_aulas"]
                detalhes["horario_semanal"] = f"Carga {contrato['percentagem']}%: {contrato['tempo_contratual']}h total / {contrato['tempo_aulas']}h aulas / {contrato['tempo_apoio']}h apoio / {contrato['tempo_preparacao']}h preparação"

        # 2. Buscar Rascunho (Se o utilizador tiver um rascunho a meio, ele ganha prioridade sobre o histórico)
        cursor.execute(
            "SELECT dados_formulario FROM rascunhos WHERE nome_docente = %s ORDER BY id DESC LIMIT 1",
            (nome,)
        )
        rascunho = cursor.fetchone()
        if rascunho and rascunho["dados_formulario"]:
            dados_form = json.loads(rascunho["dados_formulario"])
            for chave, valor in dados_form.items():
                if valor:
                    detalhes[chave] = valor

        cursor.close()
    except Exception as e:
        print(f"[!] Erro ao buscar detalhes do docente: {e}")

    return jsonify(detalhes)


@app.route('/api/cargas-horias')
def cargas_horias():
    """Retorna a lista de cargas horárias disponíveis da base de dados."""
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        cursor.execute(
            "SELECT id_carga, tempo_contratual, tempo_aulas, tempo_apoio, tempo_preparacao, percentagem FROM carga_horaria ORDER BY tempo_contratual ASC"
        )
        cargas = cursor.fetchall()
        cursor.close()
        
        # Formatar os dados para exibição
        cargas_formatadas = []
        for carga in cargas:
            # Converter Decimal para float explicitamente
            tempo_contratual = float(carga["tempo_contratual"])
            tempo_aulas = float(carga["tempo_aulas"])
            tempo_apoio = float(carga["tempo_apoio"])
            tempo_preparacao = float(carga["tempo_preparacao"])
            percentagem = float(carga["percentagem"])
            
            item = {
                "id": int(carga["id_carga"]),
                "tempo_contratual": tempo_contratual,
                "tempo_aulas": tempo_aulas,
                "tempo_apoio": tempo_apoio,
                "tempo_preparacao": tempo_preparacao,
                "percentagem": percentagem,
                "display": f"Carga {carga['id_carga']}: {tempo_contratual}h total / {tempo_aulas}h aulas / {tempo_apoio}h apoio / {tempo_preparacao}h preparação ({percentagem}%)"
            }
            cargas_formatadas.append(item)
        
        return jsonify({"cargas": cargas_formatadas})
    except Exception as e:
        print(f"[!] Erro ao buscar cargas horárias: {e}")
        return jsonify({"erro": "Erro ao buscar cargas horárias"}), 500


@app.route('/submeter-contratacao', methods=['POST'])
def submeter_contratacao():
    dados = request.get_json()
    tipo = dados.get('tipo')
    
    nome_docente_raw = dados.get('nome_docente')
    if not nome_docente_raw or not tipo:
        return jsonify({"erro": "Dados incompletos: O nome do docente é obrigatório."}), 400
    
    rascunho_id = dados.get('rascunho_id') or request.args.get('rascunho_id')
    caminho_pasta, resultado = processar_templates(dados, tipo)
    
    if caminho_pasta is None:
        return jsonify({"erro": resultado}), 400
    
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        ROOT_DIR = os.path.dirname(BASE_DIR)
        
        # ==========================================
        # PASSO 1: RESOLVER O DOCENTE
        # ==========================================
        cursor.execute("SELECT id_docente FROM docentes WHERE nome = %s", (nome_docente_raw,))
        docente_existente = cursor.fetchone()
        
        departamento = dados.get('departamento', dados.get('area_contratacao', 'matemática'))
        dept_mapping = {
            'matemática': 'matemática', 'matematica': 'matemática',
            'física': 'física', 'fisica': 'física',
            'gestão': 'gestão', 'gestao': 'gestão',
            'informática': 'matemática', 'informatica': 'matemática',
            'design': 'gestão'
        }
        departamento_valido = dept_mapping.get(departamento.strip().lower(), 'matemática')
        tipo_docente_bd = 'carreira' if tipo in ['renovacao-integral', 'renovacao-parcial'] else 'contratado'
        dados_json = json.dumps(dados)

        if not docente_existente:
            cursor.execute(
                "INSERT INTO docentes (nome, tipo_docente, departamento, dados_historico) VALUES (%s, %s, %s, %s)",
                (nome_docente_raw, tipo_docente_bd, departamento_valido, dados_json)
            )
            id_docente = cursor.lastrowid  # Captura o ID do docente acabado de criar
        else:
            id_docente = docente_existente['id_docente']
            cursor.execute(
                "UPDATE docentes SET dados_historico = %s WHERE id_docente = %s",
                (dados_json, id_docente)
            )

        # ==========================================
        # PASSO 2: DESCOBRIR A CARGA HORÁRIA
        # ==========================================
        id_da_carga = dados.get('id_carga')
        
        if not id_da_carga:
            # Fallback 1: tentar por percentagem se vier no formulário
            percentagem = dados.get('percentagem')
            if percentagem:
                cursor.execute("SELECT id_carga FROM carga_horaria WHERE percentagem = %s LIMIT 1", (percentagem,))
                resultado_carga = cursor.fetchone()
                if resultado_carga:
                    id_da_carga = resultado_carga['id_carga']
            
            # Fallback 2: tentar por total_horas_contacto (tempo_aulas)
            if not id_da_carga:
                total_horas = dados.get('total_horas_contacto')
                if total_horas:
                    cursor.execute("SELECT id_carga FROM carga_horaria WHERE tempo_aulas = %s LIMIT 1", (total_horas,))
                    resultado_carga = cursor.fetchone()
                    if resultado_carga:
                        id_da_carga = resultado_carga['id_carga']
            
            # Fallback 3 (último recurso): 100% para tempo integral, primeiro registo para parcial
            if not id_da_carga:
                if tipo == 'renovacao-integral':
                    cursor.execute("SELECT id_carga FROM carga_horaria WHERE percentagem = 100 LIMIT 1")
                    resultado_carga = cursor.fetchone()
                    if resultado_carga:
                        id_da_carga = resultado_carga['id_carga']
                
                if not id_da_carga:
                    cursor.execute("SELECT id_carga FROM carga_horaria ORDER BY id_carga ASC LIMIT 1")
                    fallback_carga = cursor.fetchone()
                    if fallback_carga:
                        id_da_carga = fallback_carga['id_carga']
                        print(f"[!] Aviso: id_carga não recebido do formulário. A usar carga de fallback: {id_da_carga}")
                    else:
                        return jsonify({"erro": "A tabela carga_horaria está vazia na BD. Insira valores de referência."}), 400

        # ==========================================
        # PASSO 3: DESCOBRIR O TEMPLATE ID
        # ==========================================
        # Mapeamento do tipo do formulário para o valor exato na tabela templates
        TEMPLATE_POR_TIPO = {
            "renovacao-integral": "Tempo Integral",
            "renovacao-parcial":  "Tempo Parcial",
            "primeira-vez":       "Tempo Parcial Edital",
        }
        tipo_template_bd = TEMPLATE_POR_TIPO.get(tipo)

        cursor.execute("SELECT id_template FROM templates WHERE tipo_contrato = %s", (tipo_template_bd,))
        resultado_template = cursor.fetchone()
        
        if resultado_template:
            id_template = resultado_template['id_template']
        else:
            # Fallback: Se não houver correspondência exata, apanha o primeiro template disponível
            cursor.execute("SELECT id_template FROM templates LIMIT 1")
            fallback_template = cursor.fetchone()
            id_template = fallback_template['id_template'] if fallback_template else None
            if not id_template:
                return jsonify({"erro": "Tabela templates vazia."}), 400

        # ==========================================
        # PASSO 4: CRIAR O CONTRATO (COM AS TRÊS CHAVES OBRIGATÓRIAS)
        # ==========================================
        cursor.execute(
            """INSERT INTO contratos 
               (docentes_id_docente, carga_horaria_id_carga, templates_id_template, data_renovacao) 
               VALUES (%s, %s, %s, %s)""",
            (id_docente, id_da_carga, id_template, datetime.now().date().isoformat())
        )
        id_novo_contrato = cursor.lastrowid

        # ==========================================
        # PASSO 5: INSERIR OS DOCUMENTOS NA BD
        # ==========================================
        ficheiros_gerados = os.listdir(caminho_pasta)
        documentos_inseridos = 0
        
        for ficheiro in ficheiros_gerados:
            caminho_ficheiro_absoluto = os.path.join(caminho_pasta, ficheiro)
            if os.path.isfile(caminho_ficheiro_absoluto):
                caminho_relativo_ficheiro = os.path.relpath(caminho_ficheiro_absoluto, ROOT_DIR).replace('\\', '/')
                
                cursor.execute(
                    """INSERT INTO documentos 
                       (nome, caminho, categoria, data_upload, contratos_id_contrato, docentes_id_docente) 
                       VALUES (%s, %s, %s, %s, %s, %s)""",
                    (ficheiro, caminho_relativo_ficheiro, 'gerado', datetime.now().isoformat(), id_novo_contrato, id_docente)
                )
                documentos_inseridos += 1
        
        if rascunho_id:
            cursor.execute("DELETE FROM rascunhos WHERE id = %s", (rascunho_id,))
        
        db.commit()
        cursor.close()
        
        return jsonify({
            "mensagem": f"Sucesso: {documentos_inseridos} documentos gerados e registados.", 
            "pasta": caminho_pasta,
            "ficheiros": ficheiros_gerados
        }), 200

    except Exception as e:
        print(f"[!] Erro ao registar na BD: {e}")
        return jsonify({"erro": f"Erro BD: {e}"}), 500


@app.route('/guardar-rascunho', methods=['POST'])
def rota_guardar_rascunho():
    dados = request.get_json()
    
    nome = dados.get('nome_docente', 'Sem Nome')
    tipo = dados.get('tipo', 'indefinido')
    
    dados_texto = json.dumps(dados)
    data_atual = datetime.now().strftime("%Y-%m-%d %H:%M")

    try:
        db = get_db()
        cursor = db.cursor()
        
        query = """
            INSERT INTO rascunhos (nome_docente, tipo_contrato, dados_formulario, data_guardado) 
            VALUES (%s, %s, %s, %s)
        """
        cursor.execute(query, (nome, tipo, dados_texto, data_atual))
        db.commit()
        
        return jsonify({"sucesso": True})
        
    except Exception as e:
        print(f"[!] Erro ao guardar rascunho: {e}")
        return jsonify({"sucesso": False, "erro": str(e)})
    finally:
        if 'cursor' in locals():
            cursor.close()
        # Removi o db.close() daqui para não fechar a ligação prematuramente. O @app.teardown_appcontext trata disso.

@app.route('/meus-rascunhos')
def ver_rascunhos():
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT * FROM rascunhos ORDER BY id DESC")
    lista_rascunhos = cursor.fetchall()
    cursor.close()
    
    return render_template('pasta_rascunhos.html', rascunhos=lista_rascunhos)

@app.route('/upload-documento', methods=['POST'])
def upload_documento():
    # Verifica se a requisição contém ficheiros
    if 'file' not in request.files:
        return jsonify({"erro": "Nenhum ficheiro foi enviado."}), 400

    ficheiros = request.files.getlist('file')
    if not ficheiros or ficheiros[0].filename == '':
        return jsonify({"erro": "Nenhum ficheiro selecionado."}), 400

    try:
        db = get_db()
        cursor = db.cursor()
        
        ROOT_DIR = os.path.dirname(BASE_DIR)
        caminho_uploads = os.path.join(ROOT_DIR, PASTA_UPLOADS)
        
        # Garantir que a pasta existe
        if not os.path.exists(caminho_uploads):
            os.makedirs(caminho_uploads)

        documentos_inseridos = 0
        
        for ficheiro in ficheiros:
            if ficheiro.filename:
                # Limpar o nome do ficheiro para evitar problemas no sistema operativo
                nome_seguro = secure_filename(ficheiro.filename)
                caminho_absoluto = os.path.join(caminho_uploads, nome_seguro)
                
                # Guardar fisicamente no disco
                ficheiro.save(caminho_absoluto)
                
                # Caminho relativo para gravar na BD
                caminho_relativo = os.path.join(PASTA_UPLOADS, nome_seguro).replace('\\', '/')
                
                # Inserir na base de dados (categoria pdf_llm conforme os teus dados)
                cursor.execute(
                    "INSERT INTO documentos (nome, caminho, categoria, data_upload) VALUES (%s, %s, %s, %s)",
                    (nome_seguro, caminho_relativo, 'pdf_llm', datetime.now().isoformat())
                )
                documentos_inseridos += 1
                
        db.commit()
        cursor.close()
        
        print(f"[OK] Upload concluído: {documentos_inseridos} ficheiros recebidos.")
        return jsonify({"mensagem": f"{documentos_inseridos} ficheiros guardados com sucesso!"}), 200

    except Exception as e:
        print(f"[!] Erro no upload: {e}")
        return jsonify({"erro": f"Erro interno ao processar o upload: {e}"}), 500

@app.route('/chat', methods=['POST'])
def chat():
    dados = request.json
    msg = dados.get('message', '')
    if not msg: return jsonify({"reply": "Escreve uma mensagem."}), 400
    if not ia_pronta: return jsonify({"reply": "IA a carregar..."}), 200
    try:
        res, _ = responder_pergunta(msg, global_retriever, global_llm, global_sql_agent)
        return jsonify({"reply": res}), 200
    except Exception as e:
        return jsonify({"reply": f"Erro: {e}"}), 500

@app.route('/ver_documento/<int:doc_id>')
def ver_documento(doc_id):
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT nome, caminho FROM documentos WHERE id = %s", (doc_id,))
    doc = cursor.fetchone()
    cursor.close()
    
    if doc:
        ROOT_DIR = os.path.dirname(BASE_DIR)
        caminho_abs = os.path.join(ROOT_DIR, doc['caminho'])
        if os.path.exists(caminho_abs): 
            return send_file(caminho_abs)
    return "Não encontrado", 404

@app.route('/download-documento/<int:doc_id>')
def download_documento(doc_id):
    """Força o download do documento (sem abrir no browser)."""
    db = get_db()
    cursor = db.cursor(dictionary=True)
    cursor.execute("SELECT nome, caminho FROM documentos WHERE id = %s", (doc_id,))
    doc = cursor.fetchone()
    cursor.close()
    
    if doc:
        ROOT_DIR = os.path.dirname(BASE_DIR)
        caminho_abs = os.path.join(ROOT_DIR, doc['caminho'])
        if os.path.exists(caminho_abs):
            return send_file(caminho_abs, as_attachment=True, download_name=doc['nome'])
    return "Não encontrado", 404

# --- ROTAS DE ELIMINAÇÃO ---
@app.route('/api/apagar-todos-rascunhos', methods=['DELETE'])
def apagar_todos_rascunhos():
    """Apaga todos os rascunhos da base de dados."""
    try:
        db = get_db()
        cursor = db.cursor()
        cursor.execute("DELETE FROM rascunhos")
        db.commit()
        cursor.close()
        
        print("[OK] Todos os rascunhos foram apagados.")
        return jsonify({"sucesso": True, "mensagem": "Todos os rascunhos foram apagados com sucesso."}), 200
    except Exception as e:
        print(f"[!] Erro ao apagar rascunhos: {e}")
        return jsonify({"sucesso": False, "erro": str(e)}), 500

@app.route('/api/apagar-rascunho/<int:rascunho_id>', methods=['DELETE'])
def apagar_rascunho(rascunho_id):
    """Apaga um rascunho específico da base de dados."""
    try:
        db = get_db()
        cursor = db.cursor()
        cursor.execute("DELETE FROM rascunhos WHERE id = %s", (rascunho_id,))
        db.commit()
        cursor.close()
        
        print(f"[OK] Rascunho {rascunho_id} foi apagado.")
        return jsonify({"sucesso": True, "mensagem": "Rascunho apagado com sucesso."}), 200
    except Exception as e:
        print(f"[!] Erro ao apagar rascunho: {e}")
        return jsonify({"sucesso": False, "erro": str(e)}), 500

@app.route('/api/apagar-documento/<int:doc_id>', methods=['DELETE'])
def apagar_documento(doc_id):
    """Apaga um documento específico - INCLUINDO toda a pasta do contrato na BD e no disco."""
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # Buscar o documento
        cursor.execute("SELECT caminho FROM documentos WHERE id = %s", (doc_id,))
        doc = cursor.fetchone()
        
        if not doc:
            cursor.close()
            print(f"[!] Documento {doc_id} não encontrado na BD")
            return jsonify({"sucesso": False, "erro": "Documento não encontrado."}), 404
        
        print(f"[DEBUG] Documento encontrado: {doc['caminho']}")
        
        # Extrair o diretório do contrato (pasta pai)
        # Exemplo: "Modelos Contratuais/Modelos Gerados/Contrato_Maria_Antonia/ficheiro.docx"
        caminho_partes = doc['caminho'].replace('\\', '/').split('/')
        print(f"[DEBUG] Partes do caminho: {caminho_partes}")
        
        # Procurar a pasta do contrato (começa com "Contrato_")
        pasta_contrato = None
        indice_contrato = -1
        
        for i, parte in enumerate(caminho_partes):
            if parte.startswith('Contrato_'):
                pasta_contrato = parte
                indice_contrato = i
                print(f"[DEBUG] Pasta do contrato encontrada: {pasta_contrato} no índice {i}")
                break
        
        if pasta_contrato:
            print(f"[DEBUG] A eliminar pasta do contrato: {pasta_contrato}")
            
            # Buscar TODOS os documentos desta pasta na BD
            cursor.execute(
                "SELECT id, caminho FROM documentos WHERE caminho LIKE %s",
                (f"%{pasta_contrato}%",)
            )
            docs_na_pasta = cursor.fetchall()
            print(f"[DEBUG] Encontrados {len(docs_na_pasta)} documentos na pasta")
            
            ROOT_DIR = os.path.dirname(BASE_DIR)
            caminho_abs_pasta = os.path.join(ROOT_DIR, *caminho_partes[:indice_contrato + 1])
            print(f"[DEBUG] Caminho absoluto da pasta: {caminho_abs_pasta}")
            print(f"[DEBUG] Pasta existe? {os.path.exists(caminho_abs_pasta)}")
            
            ficheiros_eliminados = 0
            
            # Apagar ficheiros fisicamente
            if os.path.exists(caminho_abs_pasta):
                try:
                    print(f"[DEBUG] Tentando eliminar pasta: {caminho_abs_pasta}")
                    shutil.rmtree(caminho_abs_pasta)
                    ficheiros_eliminados = len(docs_na_pasta)
                    print(f"[OK] Pasta do contrato eliminada: {caminho_abs_pasta}")
                except Exception as e:
                    print(f"[!] Erro ao eliminar pasta: {e}")
                    import traceback
                    traceback.print_exc()
                    return jsonify({"sucesso": False, "erro": f"Erro ao eliminar pasta: {e}"}), 500
            else:
                print(f"[!] Pasta não existe: {caminho_abs_pasta}")
            
            # Apagar TODOS os documentos desta pasta na BD
            print(f"[DEBUG] A eliminar registos da BD para: %{pasta_contrato}%")
            cursor.execute(
                "DELETE FROM documentos WHERE caminho LIKE %s",
                (f"%{pasta_contrato}%",)
            )
            print(f"[DEBUG] Registos eliminados: {cursor.rowcount}")
        else:
            print(f"[DEBUG] Pasta Contrato_ não encontrada, eliminando ficheiro individual")
            # Se não encontrar pasta Contrato_, apenas apagar o ficheiro individual
            ROOT_DIR = os.path.dirname(BASE_DIR)
            caminho_abs = os.path.join(ROOT_DIR, doc['caminho'])
            
            if os.path.exists(caminho_abs):
                try:
                    os.remove(caminho_abs)
                    print(f"[OK] Ficheiro eliminado: {caminho_abs}")
                except Exception as e:
                    print(f"[!] Erro ao eliminar ficheiro: {e}")
                    return jsonify({"sucesso": False, "erro": f"Erro ao eliminar ficheiro: {e}"}), 500
            
            # Apagar apenas este registo da BD
            cursor.execute("DELETE FROM documentos WHERE id = %s", (doc_id,))
        
        db.commit()
        cursor.close()
        
        print(f"[OK] Documento {doc_id} e contrato associado foram apagados.")
        return jsonify({"sucesso": True, "mensagem": "Contrato e documentos associados foram apagados com sucesso."}), 200
    except Exception as e:
        print(f"[!] Erro ao apagar documento: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": str(e)}), 500

@app.route('/api/apagar-pasta/<caso>', methods=['DELETE'])
def apagar_pasta(caso):
    """Apaga todos os documentos de uma pasta/caso específico."""
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # Converter espaços em underscores (como são guardados na BD)
        caso_normalizado = caso.replace(' ', '_')
        print(f"[DEBUG] Apagando pasta: '{caso}' -> '{caso_normalizado}'")
        
        # Buscar todos os documentos associados à pasta
        cursor.execute(
            "SELECT id, caminho FROM documentos WHERE caminho LIKE %s",
            (f"%{caso_normalizado}%",)
        )
        documentos = cursor.fetchall()
        print(f"[DEBUG] Encontrados {len(documentos)} documentos na BD")
        pastas_para_apagar = set()
        
        ROOT_DIR = os.path.dirname(BASE_DIR)
        ficheiros_eliminados = 0
        
        for doc in documentos:
            caminho_abs = os.path.join(ROOT_DIR, doc['caminho'])
            print(f"[DEBUG] Tentando eliminar ficheiro: {caminho_abs}")
            
                        # Guardar a pasta do ficheiro
            pasta_do_ficheiro = os.path.dirname(caminho_abs)
            pastas_para_apagar.add(pasta_do_ficheiro)
            
            if os.path.exists(caminho_abs):
                try:
                    os.remove(caminho_abs)
                    ficheiros_eliminados += 1
                    print(f"[OK] Ficheiro eliminado: {caminho_abs}")
                except Exception as e:
                    print(f"[!] Erro ao eliminar ficheiro {caminho_abs}: {e}")
            else:
                print(f"[!] Ficheiro não existe: {caminho_abs}")
                # Eliminar todas as pastas vazias (a partir das mais profundas)
                for pasta in sorted(pastas_para_apagar, reverse=True):
                    try:
                        if os.path.exists(pasta) and not os.listdir(pasta):  # Pasta existe e está vazia
                            os.rmdir(pasta)
                            print(f"[OK] Pasta vazia eliminada: {pasta}")
                    except Exception as e:
                        print(f"[!] Erro ao eliminar pasta {pasta}: {e}")
        
        
        # Eliminar todos os registo da base de dados
        print(f"[DEBUG] A eliminar da BD com padrão: %{caso_normalizado}%")
        cursor.execute("DELETE FROM documentos WHERE caminho LIKE %s", (f"%{caso_normalizado}%",))
        print(f"[DEBUG] Registos eliminados da BD: {cursor.rowcount}")
        db.commit()
        cursor.close()
        
        print(f"[OK] Pasta '{caso}' e conteúdo eliminados ({ficheiros_eliminados} ficheiros).")
        return jsonify({"sucesso": True, "mensagem": f"Pasta eliminada com sucesso ({ficheiros_eliminados} ficheiros removidos)."}), 200
    except Exception as e:
        print(f"[!] Erro ao apagar pasta: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({"sucesso": False, "erro": str(e)}), 500
        

#esta rita está inutilizada
#@app.route('/importar-historico')
#def importar_historico():
#    return render_template('importar.html')

@app.route('/base-dados')
def explorar_bd():
    return render_template('base_dados.html')

@app.route('/api/filtros-disponiveis')
def filtros_disponiveis():
    """Retorna os valores únicos para os filtros da consulta."""
    try:
        db = get_db()
        cursor = db.cursor()
        
        filtros = {
            "departamentos": [],
            "tipos_contrato": list(TIPOS_DOCUMENTOS_PERMITIDOS),
            "anos_letivos": []
        }
        
        cursor.execute("SELECT DISTINCT departamento FROM docentes WHERE departamento IS NOT NULL AND departamento != ''")
        filtros["departamentos"] = [r[0] for r in cursor.fetchall()]
        
        # Tentar obter anos letivos dos rascunhos
        cursor.execute("SELECT DISTINCT JSON_UNQUOTE(JSON_EXTRACT(dados_formulario, '$.ano_letivo')) FROM rascunhos WHERE JSON_EXTRACT(dados_formulario, '$.ano_letivo') IS NOT NULL")
        anos = [r[0] for r in cursor.fetchall() if r[0]]
        
        # Adicionar anos letivos de contratos se existirem (através da data_renovacao como fallback)
        cursor.execute("SELECT DISTINCT YEAR(data_renovacao) FROM contratos WHERE data_renovacao IS NOT NULL")
        anos_cont = [f"{r[0]}-{r[0]+1}" for r in cursor.fetchall() if r[0]]
        
        filtros["anos_letivos"] = sorted(list(set(anos + anos_cont)), reverse=True)
        
        cursor.close()
        return jsonify(filtros)
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/consulta-docentes')
def consulta_docentes():
    """Consulta avançada de docentes com filtros."""
    nome = request.args.get('nome', '').strip()
    dept = request.args.get('departamento', '').strip()
    tipo = request.args.get('tipo_contrato', '').strip()
    ano = request.args.get('ano_letivo', '').strip()
    
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # Base da query - usando a estrutura exata fornecida
        # Nota: isJuri na BD é "sim" ou "não"
        query = "SELECT DISTINCT d.id_docente, d.nome, d.departamento, d.tipo_docente, d.isJuri FROM docentes d"
        joins = []
        where_clauses = ["1=1"]
        params = []
        
        if nome:
            where_clauses.append("d.nome LIKE %s")
            params.append(f"%{nome}%")
        
        if dept:
            where_clauses.append("d.departamento = %s")
            params.append(dept)
            
        if tipo:
            joins.append("LEFT JOIN contratos c ON d.id_docente = c.docentes_id_docente")
            joins.append("LEFT JOIN templates t ON c.templates_id_template = t.id_template")
            where_clauses.append("(t.tipo_contrato = %s OR d.dados_historico LIKE %s)")
            params.append(tipo)
            params.append(f"%{tipo}%")

        if ano:
            where_clauses.append("(d.dados_historico LIKE %s OR EXISTS (SELECT 1 FROM rascunhos r WHERE r.nome_docente = d.nome AND r.dados_formulario LIKE %s))")
            params.append(f"%{ano}%")
            params.append(f"%{ano}%")
        
        final_query = f"{query} {' '.join(joins)} WHERE {' AND '.join(where_clauses)}"
        
        cursor.execute(final_query, tuple(params))
        docentes = cursor.fetchall()
        
        cursor.close()
        return jsonify({"docentes": docentes})
    except Exception as e:
        print(f"[!] Erro na consulta: {e}")
        return jsonify({"erro": str(e)}), 500

@app.route('/api/detalhe-docente/<int:id_docente>')
def detalhe_docente(id_docente):
    """Retorna todos os dados associados a um docente usando a estrutura completa."""
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # 1. Dados Pessoais (docentes: id_docente, nome, tipo_docente, departamento, dados_historico, isJuri)
        cursor.execute("SELECT * FROM docentes WHERE id_docente = %s", (id_docente,))
        docente = cursor.fetchone()
        if not docente:
            return jsonify({"erro": "Docente não encontrado"}), 404
        
        if docente.get('dados_historico'):
            try:
                docente['historico_json'] = json.loads(docente['dados_historico'])
            except:
                docente['historico_json'] = {}
        
        # 2. Contratos e Carga Horária
        cursor.execute("""
            SELECT 
                c.id_contrato, c.numero_renovacao, c.data_renovacao,
                ch.tempo_contratual, ch.tempo_aulas, ch.tempo_apoio, ch.tempo_preparacao, ch.percentagem,
                t.tipo_contrato as tipo_template, t.caminho_ficheiro as template_path
            FROM contratos c
            LEFT JOIN carga_horaria ch ON c.carga_horaria_id_carga = ch.id_carga
            LEFT JOIN templates t ON c.templates_id_template = t.id_template
            WHERE c.docentes_id_docente = %s
            ORDER BY c.data_renovacao DESC
        """, (id_docente,))
        contratos = cursor.fetchall()

        # Mapeamento de tipo da tabela templates → label legível para o frontend
        TIPO_LEGIVEL = {
            "Tempo Integral":      "Tempo Integral",
            "Tempo Parcial":       "Tempo Parcial",
            "Tempo Parcial Edital":"Tempo Parcial (Edital)",
        }

        for c in contratos:
            for k, v in c.items():
                if hasattr(v, '__float__'):
                    c[k] = float(v)
                elif isinstance(v, (datetime, date)):
                    c[k] = v.isoformat()

            # Tipo legível (dentro do loop — um por contrato)
            c['tipo_legivel'] = TIPO_LEGIVEL.get(c.get('tipo_template'), c.get('tipo_template') or 'Contrato')

            # Ano letivo: buscar no caminho do documento associado a este contrato
            cursor.execute(
                "SELECT caminho FROM documentos WHERE contratos_id_contrato = %s LIMIT 1",
                (c['id_contrato'],)
            )
            doc = cursor.fetchone()
            if doc:
                # O caminho tem a estrutura: .../Contrato_Nome/2025-2026/ficheiro.docx
                partes = doc['caminho'].replace('\\', '/').split('/')
                ano = next((p for p in partes if '-' in p and len(p) == 9 and p[:4].isdigit()), None)
                c['ano_letivo'] = ano
            else:
                c['ano_letivo'] = None

        # 3. Documentos (documentos: id, nome, caminho, categoria, data_upload, versao_contrato, contratos_id_contrato, docentes_id_docente)
        cursor.execute("""
            SELECT id, nome, caminho, categoria, data_upload, versao_contrato, contratos_id_contrato 
            FROM documentos 
            WHERE docentes_id_docente = %s
        """, (id_docente,))
        documentos = cursor.fetchall()
        for d in documentos:
            if isinstance(d.get('data_upload'), (datetime, date)):
                d['data_upload'] = d['data_upload'].isoformat()

        # 4. Rascunhos
        cursor.execute("SELECT id, tipo_contrato, data_guardado, dados_formulario FROM rascunhos WHERE nome_docente = %s", (docente['nome'],))
        rascunhos = cursor.fetchall()
        for r in rascunhos:
            if r.get('dados_formulario'):
                try:
                    r['dados_formulario'] = json.loads(r['dados_formulario'])
                except:
                    pass

        cursor.close()
        return jsonify({
            "docente": docente,
            "contratos": contratos,
            "documentos": documentos,
            "rascunhos": rascunhos
        })
    except Exception as e:
        print(f"[!] Erro no detalhe: {e}")
        return jsonify({"erro": str(e)}), 500

@app.route('/api/procurar-docente')
def procurar_docente():
    nome = request.args.get('q', '').strip()
    if not nome:
        return jsonify({"erro": "Insira um nome"}), 400
    
    resultado = {"docente": None, "rascunhos": [], "documentos": []}

    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        #1. Procurar perfil na tabela docentes
        cursor.execute("SELECT * FROM docentes WHERE nome LIKE %s", (f"%{nome}%",))
        docente = cursor.fetchone()

        if docente:
            #Se houver histórico guardado em JSON, converte para enviar bonito
            if docente.get('dados_historico'):
                docente['historico:json'] = json.loads(docente['dados_historico'])
            resultado['docente'] = docente

            #2. procurar rascunhos associados
            cursor.execute("SELECT id, tipo_contrato, data_guardado FROM rascunhos WHERE nome_docente LIKE %s", (f"%{nome}%",))
            resultado['rascunhos'] = cursor.fetchall()

            #3. procurar documentos gerados (pela pasta Contrato_Nome)
            nome_pasta = f"Contrato_{docente['nome'].replace(' ', '_')}"
            cursor.execute("SELECT id, nome, categoria, data_upload, caminho FROM documentos WHERE caminho LIKE %s", (f"%{nome_pasta}%",))
            resultado['documentos'] = cursor.fetchall()

        cursor.close()
        return jsonify(resultado)
    except Exception as e:
        print(f"[!] Erro na pesquisa: {e}")
        return jsonify({"erro": str(e)}), 500
    

@app.route('/api/lista-tabelas')
def lista_tabelas():
    try:
        db = get_db()
        cursor = db.cursor()
        cursor.execute("SHOW TABLES")
        # O cursor devolve tuplos como [('docentes',), ('contratos',)], transformamos numa lista simples
        tabelas = [t[0] for t in cursor.fetchall()]
        cursor.close()
        return jsonify({"tabelas": tabelas})
    except Exception as e:
        return jsonify({"erro": str(e)}), 500

@app.route('/api/dados-tabela/<table_name>')
def dados_tabela(table_name):
    try:
        db = get_db()
        # Verificação de segurança dinâmica: a tabela existe realmente na BD?
        cursor = db.cursor()
        cursor.execute("SHOW TABLES")
        tabelas_existentes = [t[0] for t in cursor.fetchall()]
        cursor.close()

        if table_name not in tabelas_existentes:
            return jsonify({"erro": "Tabela não encontrada"}), 404
        
        # Agora sim, fazemos a consulta
        cursor = db.cursor(dictionary=True)
        cursor.execute(f"SELECT * FROM `{table_name}` LIMIT 100")
        dados = cursor.fetchall()
        cursor.close()
        return jsonify({"dados": dados})
    except Exception as e:
        return jsonify({"erro": str(e)}), 500
    
@app.route('/api/update-cell', methods=['POST'])
def update_cell():
    data = request.json
    table = data.get('table')
    coluna = data.get('column')
    valor = data.get('value')
    # O identificador é sempre a primeira coluna da tabela (assumindo que é a PK)
    id_col = data.get('id_col') 
    id_val = data.get('id_val')

    try:
        db = get_db()
        cursor = db.cursor()
        # Query parametrizada para evitar SQL Injection
        query = f"UPDATE `{table}` SET `{coluna}` = %s WHERE `{id_col}` = %s"
        cursor.execute(query, (valor, id_val))
        db.commit()
        cursor.close()
        return jsonify({"status": "sucesso"})
    except Exception as e:
        return jsonify({"status": "erro", "message": str(e)}), 500
    
@app.route('/api/adicionar-linha', methods=['POST'])
def adicionar_linha():
    data = request.json
    table = data.get('table')
    # Validar se a tabela existe
    db = get_db()
    cursor = db.cursor()
    cursor.execute(f"INSERT INTO `{table}` () VALUES ()") # Cria uma linha com valores default
    db.commit()
    cursor.close()
    return jsonify({"status": "sucesso"})

@app.route('/api/apagar-linha', methods=['POST'])
def apagar_linha():
    data = request.json
    table = data.get('table')
    id_col = data.get('id_col')
    id_val = data.get('id_val')
    
    db = get_db()
    cursor = db.cursor()
    cursor.execute(f"DELETE FROM `{table}` WHERE `{id_col}` = %s", (id_val,))
    db.commit()
    cursor.close()
    return jsonify({"status": "sucesso"})

@app.route('/api/listar-pastas-templates')
def listar_pastas_templates():
    """Devolve as pastas únicas de templates disponíveis, agrupadas por categoria."""
    try:
        db = get_db()
        cursor = db.cursor(dictionary=True)
        cursor.execute("""
            SELECT DISTINCT categoria, caminho 
            FROM documentos 
            WHERE categoria IN ('tempo integral anual', 'tempo parcial semestral', 'tempo parcial edital')
            AND caminho NOT LIKE '%Modelos Gerados%'
        """)
        docs = cursor.fetchall()
        cursor.close()

        pastas = {}
        vistos = set()

        for doc in docs:
            cat = doc['categoria']
            caminho_pasta = os.path.dirname(doc['caminho'])
            if not caminho_pasta or caminho_pasta in vistos:
                continue
            vistos.add(caminho_pasta)

            if cat not in pastas:
                pastas[cat] = []
            pastas[cat].append({
                "label": caminho_pasta,
                "caminho": caminho_pasta,
                "categoria": cat
            })

        return jsonify({"pastas": pastas})
    except Exception as e:
        print(f"[!] Erro ao listar pastas de templates: {e}")
        return jsonify({"erro": str(e)}), 500


@app.route('/api/upload-template', methods=['POST'])
def upload_template():
    """Faz upload de um ficheiro para uma pasta de templates existente e regista na BD."""
    if 'file' not in request.files:
        return jsonify({"erro": "Nenhum ficheiro enviado."}), 400

    ficheiro = request.files['file']
    caminho_destino = request.form.get('pasta', '').strip()
    categoria = request.form.get('categoria', '').strip()

    if not ficheiro.filename or not caminho_destino or not categoria:
        return jsonify({"erro": "Dados incompletos: ficheiro, pasta e categoria são obrigatórios."}), 400

    # Validar categoria para evitar injeção
    if categoria not in TIPOS_DOCUMENTOS_PERMITIDOS:
        return jsonify({"erro": "Categoria inválida."}), 400

    try:
        ROOT_DIR = os.path.dirname(BASE_DIR)
        pasta_abs = os.path.join(ROOT_DIR, caminho_destino)

        if not os.path.exists(pasta_abs):
            return jsonify({"erro": f"Pasta de destino não existe: {caminho_destino}"}), 400

        nome_seguro = secure_filename(ficheiro.filename)
        caminho_abs_ficheiro = os.path.join(pasta_abs, nome_seguro)
        ficheiro.save(caminho_abs_ficheiro)

        caminho_relativo = os.path.join(caminho_destino, nome_seguro).replace('\\', '/')

        db = get_db()
        cursor = db.cursor()
        cursor.execute(
            "INSERT INTO documentos (nome, caminho, categoria, data_upload) VALUES (%s, %s, %s, %s)",
            (nome_seguro, caminho_relativo, categoria, datetime.now().isoformat())
        )
        db.commit()
        cursor.close()

        print(f"[OK] Template adicionado: {caminho_relativo} (categoria: {categoria})")
        return jsonify({"mensagem": f"'{nome_seguro}' adicionado com sucesso!", "nome": nome_seguro}), 200
    except Exception as e:
        print(f"[!] Erro no upload de template: {e}")
        return jsonify({"erro": str(e)}), 500


@app.route('/api/renomear-documento/<int:doc_id>', methods=['POST'])
def renomear_documento(doc_id):
    """Renomeia um documento."""
    try:
        dados = request.get_json()
        novo_nome = dados.get('novo_nome', '').strip()
        
        if not novo_nome:
            return jsonify({"sucesso": False, "erro": "Nome não pode estar vazio."}), 400
        
        db = get_db()
        cursor = db.cursor(dictionary=True)
        
        # Buscar o documento atual
        cursor.execute("SELECT id, nome, caminho FROM documentos WHERE id = %s", (doc_id,))
        doc = cursor.fetchone()
        
        if not doc:
            cursor.close()
            return jsonify({"sucesso": False, "erro": "Documento não encontrado."}), 404
        
        nome_antigo = doc['nome']
        
        # Se o nome é igual, não fazer nada
        if nome_antigo == novo_nome:
            cursor.close()
            return jsonify({"sucesso": True, "mensagem": "Documento não foi alterado."}), 200
        
        # Atualizar na BD
        cursor.execute(
            "UPDATE documentos SET nome = %s WHERE id = %s",
            (novo_nome, doc_id)
        )
        db.commit()
        cursor.close()
        
        print(f"[OK] Documento renomeado: '{nome_antigo}' -> '{novo_nome}'")
        return jsonify({"sucesso": True, "mensagem": "Documento renomeado com sucesso."}), 200
    except Exception as e:
        print(f"[!] Erro ao renomear documento: {e}")
        return jsonify({"sucesso": False, "erro": str(e)}), 500

if __name__ == '__main__':
    app.run(host="localhost", port=5000, debug=True)